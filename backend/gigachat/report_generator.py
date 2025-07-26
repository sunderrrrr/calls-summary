from docx import Document
from fpdf import FPDF

def save_report(summary: str, path: str, fmt: str):
    if fmt == "txt":
        with open(path, "w", encoding="utf-8") as f:
            f.write(summary)
    elif fmt == "docx":
        doc = Document()
        doc.add_heading("Summary", 0)
        doc.add_paragraph(summary)
        doc.save(path)
    elif fmt == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in summary.splitlines():
            pdf.cell(0, 10, txt=line, ln=1)
        pdf.output(path)
    else:
        raise ValueError("Unsupported format")